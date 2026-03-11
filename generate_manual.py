#!/usr/bin/env python3
"""鳥取市テニス協会 ドロー会議システム 操作マニュアル Word生成"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ---- スタイル設定 ----
style = doc.styles['Normal']
font = style.font
font.name = 'Yu Gothic'
font.size = Pt(10.5)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Yu Gothic'
    h.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

# ページ設定 A4
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# ---- 表紙 ----
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('鳥取市テニス協会')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2b, 0x7e, 0xc6)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ドロー会議システム')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('操作マニュアル')
run.font.size = Pt(18)
run.bold = True

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('URL: https://tcta-tottori.github.io/tottori-tennis-draw/')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x2b, 0x7e, 0xc6)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('令和7年3月 作成')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ---- 目次 ----
doc.add_heading('目次', level=1)

toc_items = [
    ('1.', 'システム概要', '3'),
    ('2.', 'アクセス方法', '3'),
    ('3.', '画面構成', '3'),
    ('4.', '操作手順', '4'),
    ('  4.1', 'データ読込', '4'),
    ('  4.2', '選手一覧', '5'),
    ('  4.3', 'OCR入力', '5'),
    ('  4.4', 'エントリー', '6'),
    ('  4.5', '種目別確認', '7'),
    ('  4.6', '抽選', '7'),
    ('  4.7', 'ドロー表', '8'),
    ('5.', 'JTAシードルール', '9'),
    ('6.', '対応種目一覧', '10'),
    ('7.', 'トラブルシューティング', '11'),
    ('8.', '注意事項', '11'),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(11)

doc.add_page_break()

# ---- 1. システム概要 ----
doc.add_heading('1. システム概要', level=1)

doc.add_paragraph(
    '本システムは、鳥取市テニス協会が主催するテニス大会のドロー（トーナメント組み合わせ）を'
    '効率的に作成するためのWebアプリケーションです。'
)

doc.add_heading('主な機能', level=2)
features = [
    'ランキングデータの自動読み込み（Google スプレッドシート / Excelファイル）',
    'OCR（画像認識）による申込用紙の自動読み取り',
    'エントリー管理（手動追加・編集・削除）',
    'JTAルールに基づくシード自動決定',
    '手動ドロー抽選（視覚的なトーナメント表でリアルタイム確認）',
    'ドロー表のExcel / CSV出力',
    'ドロー表の印刷',
    'シングルス・ダブルス全17種目対応',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

# ---- 2. アクセス方法 ----
doc.add_heading('2. アクセス方法', level=1)

doc.add_paragraph('以下のURLにアクセスしてください。PCでもスマートフォンでも利用可能です。')

p = doc.add_paragraph()
run = p.add_run('https://tcta-tottori.github.io/tottori-tennis-draw/')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x2b, 0x7e, 0xc6)
run.bold = True

doc.add_paragraph()
notes = [
    'インターネット接続が必要です。',
    'ログイン不要 — URLを知っている人は誰でも利用できます。',
    'エントリーデータはブラウザのローカルストレージに保存されます。',
    '推奨ブラウザ: Google Chrome、Safari、Microsoft Edge',
]
for n in notes:
    doc.add_paragraph(n, style='List Bullet')

# ---- 3. 画面構成 ----
doc.add_heading('3. 画面構成', level=1)

doc.add_paragraph('画面上部のタブで各機能に切り替えます。')

table = doc.add_table(rows=9, cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = [('タブ名', '機能概要')]
data = [
    ('データ読込', 'ランキングデータとふりがなデータを読み込みます'),
    ('選手一覧', '読み込んだランキングデータを一覧表示します'),
    ('OCR入力', '申込用紙を撮影・画像認識してエントリー登録します'),
    ('エントリー', 'エントリー選手の管理（追加・編集・削除）を行います'),
    ('種目別確認', '種目ごとのエントリー状況・ドロー情報を確認します'),
    ('抽選', 'シード自動決定＋手動ドロー抽選を行います'),
    ('ドロー表', '確定したトーナメント表の表示・出力・印刷を行います'),
    ('マニュアル', '操作説明とJTAルールを確認できます'),
]

for i, (h1, h2) in enumerate(headers + data):
    row = table.rows[i]
    row.cells[0].text = h1
    row.cells[1].text = h2
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_page_break()

# ---- 4. 操作手順 ----
doc.add_heading('4. 操作手順', level=1)

doc.add_paragraph('大会のドロー作成は、以下の流れで行います。')

flow = [
    'データ読込 → ランキングデータを読み込む',
    '選手一覧 → 読み込んだデータを確認する',
    'エントリー → 出場選手を登録する（手動追加 or OCR入力）',
    '種目別確認 → 各種目のエントリー状況を確認する',
    '抽選 → 手動でドロー抽選を実行する',
    'ドロー表 → 完成したトーナメント表を確認・出力・印刷する',
]
for i, f in enumerate(flow, 1):
    doc.add_paragraph(f'Step {i}: {f}', style='List Number')

# ---- 4.1 データ読込 ----
doc.add_heading('4.1 データ読込', level=2)

doc.add_paragraph(
    '起動時にデフォルトのGoogle スプレッドシートから自動的にランキングデータを読み込みます。'
    '一度読み込んだデータはブラウザにバックアップされ、次回以降は即座に利用可能です。'
)

doc.add_heading('Google スプレッドシートから読み込み', level=3)
steps = [
    'スプレッドシートの共有設定を「リンクを知っている全員が閲覧可」に設定する',
    'スプレッドシートのURLを入力欄に貼り付ける',
    '「読込」ボタンをクリックする',
    '読込完了後、ステータスが緑色（✓）に変わる',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('ローカルファイルから読み込み', level=3)
steps = [
    'rank.xlsx（ランキングデータ）をドラッグ＆ドロップまたはクリックして選択',
    'ふりがなデータも同様にアップロード',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

p = doc.add_paragraph()
run = p.add_run('※ ')
run.bold = True
run = p.add_run('スプレッドシートの全17シート（シングルス8種目＋ダブルス9種目）を自動で読み込みます。')

# ---- 4.2 選手一覧 ----
doc.add_heading('4.2 選手一覧', level=2)

doc.add_paragraph(
    '読み込んだランキングデータの全選手を一覧表示します。'
)

features = [
    'シングルス／ダブルスの切り替え — 画面上部のトグルボタンで表示カテゴリを切り替え',
    '種目タブ — 各種目ごとに絞り込み表示',
    '検索 — 氏名・ふりがな・所属で検索可能',
    'エントリーボタン — 各選手の行にある「エントリー」ボタンで即座に登録',
    'エントリー済み表示 — 既にエントリー済みの選手はグレー表示で「登録済」バッジ表示',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

# ---- 4.3 OCR入力 ----
doc.add_heading('4.3 OCR入力', level=2)

doc.add_paragraph(
    '大会申込用紙の写真を撮影し、画像認識（OCR）で選手名を自動認識してエントリー登録できます。'
)

steps = [
    '「カメラ撮影」または「ファイル選択」をクリック',
    '申込用紙の画像を選択または撮影する',
    '「OCR実行」ボタンをクリックして認識開始',
    '認識結果の一覧が表示される — 氏名候補がドロップダウンで表示される',
    '正しい氏名・所属・種目を確認／修正する',
    '登録したい行にチェックを入れ、「確認して登録」をクリック',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

p = doc.add_paragraph()
run = p.add_run('ヒント: ')
run.bold = True
run = p.add_run('鮮明でまっすぐに撮影した画像ほど認識精度が向上します。')

doc.add_page_break()

# ---- 4.4 エントリー ----
doc.add_heading('4.4 エントリー', level=2)

doc.add_paragraph(
    '大会に出場する選手のエントリー管理を行う画面です。'
)

doc.add_heading('エントリー情報バー', level=3)
doc.add_paragraph(
    '種目フィルターの下に、現在の種目のエントリー数・ドローサイズ・BYE数・シード数が自動表示されます。'
    'シード対象選手も名前付きで表示されるため、常に最新の状態が確認できます。'
)

doc.add_heading('手動追加', level=3)
steps = [
    '「手動追加」ボタンをクリック',
    '氏名を入力 — 入力中にランキングデータから候補が自動表示される',
    '候補をクリックすると所属・種目・ポイントが自動入力される',
    '必要に応じて情報を修正し「保存」をクリック',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

p = doc.add_paragraph()
run = p.add_run('※ ')
run.bold = True
run = p.add_run(
    'エントリーが1種目のみの場合、種目は自動選択されます。'
    '所属が空欄の場合は自動的に「フリー」が設定されます。'
)

doc.add_heading('エントリーの並び順', level=3)
doc.add_paragraph('エントリー一覧はランキングポイント降順（高い順）で自動ソートされます。')

doc.add_heading('データのバックアップ', level=3)
backup_items = [
    '「エクスポート」— エントリーデータをJSONファイルとしてダウンロード',
    '「インポート」— エクスポートしたJSONファイルからデータを復元',
]
for b in backup_items:
    doc.add_paragraph(b, style='List Bullet')

# ---- 4.5 種目別確認 ----
doc.add_heading('4.5 種目別確認', level=2)

doc.add_paragraph(
    '種目ごとのエントリー状況を確認する画面です。'
    'エントリーがある種目がタブとして表示され、最初の種目が自動選択されます。'
)

doc.add_paragraph('各種目について以下の情報が確認できます:')
items = [
    'エントリー数',
    'ドローサイズ（自動計算）',
    'シード数（JTAルールに基づいて自動計算）',
    '選手一覧（ポイント降順）',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ---- 4.6 抽選 ----
doc.add_heading('4.6 抽選', level=2)

doc.add_paragraph(
    '手動でドロー抽選を行う画面です。種目を選択すると、シードが自動計算され、'
    'すぐに手動配置が開始できます。'
)

doc.add_heading('抽選の流れ', level=3)
steps = [
    '種目を選択（最初の種目が自動選択される）',
    'シード情報バーでシード選手・ドロー情報を確認',
    'シード選手とBYEは自動的に正しい位置に配置される',
    '未配置選手リストから選手をクリックして選択（青くハイライトされる）',
    'トーナメント表または配置テーブルの空きポジションをクリックして配置',
    '全選手を配置したら「ドロー確定」をクリック',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('操作のポイント', level=3)
tips = [
    '選手チップをクリックで選択 → 空きスロットをクリックで配置',
    '配置済みの非シード選手は「取消」ボタンで未配置に戻せる',
    '「リセット」ボタンでシード・BYE以外の配置を全てやり直せる',
    'トーナメント表プレビューがリアルタイムで更新される',
    '「全種目一括抽選」ボタンで全種目を自動抽選することも可能',
]
for t in tips:
    doc.add_paragraph(t, style='List Bullet')

doc.add_page_break()

# ---- 4.7 ドロー表 ----
doc.add_heading('4.7 ドロー表', level=2)

doc.add_paragraph(
    '確定したトーナメント表を表示・出力する画面です。'
    '確定済みのドローがある場合、最初の種目が自動表示されます。'
)

doc.add_heading('表示形式', level=3)
doc.add_paragraph(
    'トーナメント表はSVGで描画され、左右の山（トップハーフ・ボトムハーフ）に分かれて表示されます。'
    '各選手名の横に所属が（）付きで表示されます。'
)

doc.add_heading('出力機能', level=3)
outputs = [
    'Excel出力 — トーナメント罫線付きのExcelファイルをダウンロード（エントリーリストシート付き）',
    'CSV出力 — UTF-8のCSVファイルをダウンロード',
    '印刷 — ブラウザの印刷機能でトーナメント表を印刷（A3横推奨）',
]
for o in outputs:
    doc.add_paragraph(o, style='List Bullet')

# ---- 5. JTAシードルール ----
doc.add_heading('5. JTAシードルール', level=1)

doc.add_paragraph(
    '本システムは、JTA（日本テニス協会）公式トーナメント競技関連規則に基づいてシードを決定します。'
)

doc.add_heading('シード数の決定', level=2)

table = doc.add_table(rows=5, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
seed_data = [
    ('出場選手数', 'ドローサイズ', 'シード数'),
    ('5〜8名', '8ドロー', '2シード'),
    ('9〜16名', '16ドロー', '4シード'),
    ('17〜32名', '32ドロー', '4シード'),
    ('33名以上', '64ドロー以上', '8シード'),
]
for i, (c1, c2, c3) in enumerate(seed_data):
    row = table.rows[i]
    row.cells[0].text = c1
    row.cells[1].text = c2
    row.cells[2].text = c3
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

doc.add_heading('シード配置ルール', level=2)

rules = [
    'シード1: ドロー最上段（ライン1）に固定配置',
    'シード2: ドロー最下段（最終ライン）に固定配置',
    'シード3・4: 規定の2つの候補位置から抽選で決定',
    'シード5〜8: 規定の4つの候補位置から抽選で決定',
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('BYE配置ルール', level=2)

bye_rules = [
    'シード順位が高い選手の対戦相手位置にBYEを優先配置',
    '残りのBYEはドロー表の端（上端・下端交互）から配置',
    'BYEの対戦相手は1回戦不戦勝（自動的に2回戦進出）',
]
for r in bye_rules:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ---- 6. 対応種目一覧 ----
doc.add_heading('6. 対応種目一覧', level=1)

doc.add_paragraph('本システムは以下の17種目に対応しています。')

doc.add_heading('シングルス（8種目）', level=2)

table = doc.add_table(rows=9, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
singles = [
    ('コード', '種目名', '略称'),
    ('ms', '一般男子シングルス', '男子S'),
    ('ls', '一般女子シングルス', '女子S'),
    ('m35s', '男子35歳以上シングルス', '男子35S'),
    ('m45s', '男子45歳以上シングルス', '男子45S'),
    ('m55s', '男子55歳以上シングルス', '男子55S'),
    ('m65s', '男子65歳以上シングルス', '男子65S'),
    ('mbs', '男子B級シングルス', '男子BS'),
    ('lbs', '女子B級シングルス', '女子BS'),
]
for i, (c1, c2, c3) in enumerate(singles):
    row = table.rows[i]
    row.cells[0].text = c1
    row.cells[1].text = c2
    row.cells[2].text = c3
    if i == 0:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()

doc.add_heading('ダブルス（9種目）', level=2)

table = doc.add_table(rows=10, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
doubles = [
    ('コード', '種目名', '略称'),
    ('md', '一般男子ダブルス', '男子D'),
    ('ld', '一般女子ダブルス', '女子D'),
    ('m45d', '男子45歳以上ダブルス', '男子45D'),
    ('m55d', '男子55歳以上ダブルス', '男子55D'),
    ('m65d', '男子65歳以上ダブルス', '男子65D'),
    ('l45d', '女子45歳以上ダブルス', '女子45D'),
    ('l55d', '女子55歳以上ダブルス', '女子55D'),
    ('mbd', '男子B級ダブルス', '男子BD'),
    ('lbd', '女子B級ダブルス', '女子BD'),
]
for i, (c1, c2, c3) in enumerate(doubles):
    row = table.rows[i]
    row.cells[0].text = c1
    row.cells[1].text = c2
    row.cells[2].text = c3
    if i == 0:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_page_break()

# ---- 7. トラブルシューティング ----
doc.add_heading('7. トラブルシューティング', level=1)

troubles = [
    (
        'データが読み込めない',
        [
            'スプレッドシートの共有設定が「リンクを知っている全員が閲覧可」になっているか確認',
            'URLが正しいか確認',
            'インターネット接続を確認',
            'ブラウザのコンソール（F12キー）でエラーメッセージを確認',
        ]
    ),
    (
        'OCR認識精度が低い',
        [
            'できるだけ鮮明な画像を使用する',
            '申込用紙はまっすぐ（傾きなく）撮影する',
            '影や反射を避ける',
            '認識結果の氏名候補ドロップダウンから正しい選手を選択できます',
        ]
    ),
    (
        'エントリーデータが消えた',
        [
            'エントリーデータはブラウザのlocalStorageに保存されています',
            'ブラウザのデータをクリアすると消えるため、事前に「エクスポート」でバックアップしてください',
            '「インポート」でバックアップファイルから復元できます',
        ]
    ),
    (
        'ドロー表の印刷がうまくいかない',
        [
            'ブラウザのズーム倍率を100%に設定',
            '用紙サイズ: A3横推奨（32ドロー以上の場合）',
            '印刷プレビューで確認してから印刷',
            'Excel出力してExcelから印刷することも可能',
        ]
    ),
    (
        '別のPCで続きの作業をしたい',
        [
            'エントリー画面の「エクスポート」でJSONファイルを保存',
            '別のPCでシステムにアクセスし「インポート」でJSONファイルを読み込む',
            'ランキングデータは自動的にスプレッドシートから再読み込みされます',
        ]
    ),
]

for title, items in troubles:
    doc.add_heading(title, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ---- 8. 注意事項 ----
doc.add_heading('8. 注意事項', level=1)

notes = [
    'エントリーデータはブラウザのローカルストレージに保存されます。ブラウザのデータを削除すると失われるため、重要なデータは必ずエクスポートしてバックアップしてください。',
    'ランキングデータは一度読み込むとバックアップされますが、スプレッドシートが更新された場合は再読み込みが必要です。',
    '複数人で同時に操作する場合、各端末のデータは独立しています。一つの端末で作業し、完成したドローをExcel出力で共有することを推奨します。',
    'ドロー確定後も、抽選画面で同じ種目のドローをやり直すことが可能です（前の確定データは上書きされます）。',
    '印刷時は、32ドロー以上の場合はA3横向きでの印刷を推奨します。',
]

for n in notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('― 以上 ―')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ---- 保存 ----
output_path = os.path.join(os.path.dirname(__file__), 'ドロー会議システム_操作マニュアル.docx')
doc.save(output_path)
print(f'マニュアルを生成しました: {output_path}')
